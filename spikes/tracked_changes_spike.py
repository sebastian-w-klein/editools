"""Spike v2: corrected architecture.
Fixes found by v1:
  1. runs must be walked recursively; runs inside w:del are NOT part of current text,
     runs inside w:ins ARE.
  2. italic detection must honour w:val="0"/"false", not just <w:i> presence.
  3. edits must be applied in DESCENDING offset order so earlier offsets stay valid.
"""
import copy, re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
AUTHOR, STAMP = "Index Checker", "2026-08-22T00:00:00Z"
_id = [1000]
def nid():
    _id[0] += 1; return str(_id[0])

def new_el(tag, **attrs):
    el = etree.Element(qn(tag))
    for k, v in attrs.items(): el.set(qn(k), v)
    return el

def is_on(rPr, tag):
    """True if a toggle property is genuinely on (absent w:val means on)."""
    if rPr is None: return False
    el = rPr.find(W + tag)
    if el is None: return False
    v = el.get(qn('w:val'))
    return v not in ('0', 'false', 'off')

def live_runs(p):
    """Runs that contribute to the CURRENT text: skip anything inside w:del."""
    out = []
    for r in p.iter(W + 'r'):
        if any(a.tag == W + 'del' for a in r.iterancestors()):
            continue
        out.append(r)
    return out

def text_and_fmt(p):
    """(text, [italic_bool per char]) for the current text of the paragraph."""
    chars, ital = [], []
    for r in live_runs(p):
        t = r.find(W + 't')
        if t is None or not t.text: continue
        on = is_on(r.find(W + 'rPr'), 'i')
        chars.append(t.text); ital.extend([on] * len(t.text))
    return ''.join(chars), ital

def split_run(run, offset):
    t = run.find(W + 't'); text = t.text or ""
    if offset <= 0 or offset >= len(text): return None
    new = copy.deepcopy(run)
    t.text = text[:offset]; t.set(qn('xml:space'), 'preserve')
    nt = new.find(W + 't'); nt.text = text[offset:]; nt.set(qn('xml:space'), 'preserve')
    run.addnext(new)
    return new

def isolate(p, start, end):
    """Runs covering exactly chars [start,end) of the current text."""
    pos, targets = 0, []
    for r in live_runs(p):
        t = r.find(W + 't')
        if t is None or not t.text: continue
        length = len(t.text); r_start, r_end = pos, pos + length; pos = r_end
        if r_end <= start or r_start >= end: continue
        cur = r
        if r_start < start:
            _h = split_run(cur, start - r_start)
            cur = cur if _h is None else _h
            r_start = start
        ct = cur.find(W + 't')
        if r_start + len(ct.text) > end:
            split_run(cur, end - r_start)
        targets.append(cur)
    return targets

def track_delete(p, start, end):
    targets = isolate(p, start, end)
    if not targets: return
    d = new_el('w:del', **{'w:id': nid(), 'w:author': AUTHOR, 'w:date': STAMP})
    targets[0].addprevious(d)
    for r in targets:
        t = r.find(W + 't'); dt = new_el('w:delText')
        dt.text = t.text; dt.set(qn('xml:space'), 'preserve')
        r.remove(t); r.append(dt); d.append(r)

def track_insert(p, at, text, italic=False, model_run=None):
    ins = new_el('w:ins', **{'w:id': nid(), 'w:author': AUTHOR, 'w:date': STAMP})
    r = new_el('w:r')
    if italic:
        rPr = new_el('w:rPr'); rPr.append(new_el('w:i')); r.append(rPr)
    t = new_el('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    r.append(t); ins.append(r)
    tgt = isolate(p, at, at + 1)
    if tgt: tgt[0].addprevious(ins)
    else:
        live = live_runs(p)
        (live[-1].getparent() if live else p).append(ins) if not live else live[-1].addnext(ins)

def highlight(p, start, end, color="yellow"):
    for r in isolate(p, start, end):
        rPr = r.find(W + 'rPr')
        if rPr is None: rPr = new_el('w:rPr'); r.insert(0, rPr)
        for old in rPr.findall(W + 'highlight'): rPr.remove(old)
        rPr.append(new_el('w:highlight', **{'w:val': color}))

def track_italicize(p, start, end):
    for r in isolate(p, start, end):
        rPr = r.find(W + 'rPr')
        old = copy.deepcopy(rPr) if rPr is not None else None
        if rPr is None: rPr = new_el('w:rPr'); r.insert(0, rPr)
        for e in rPr.findall(W + 'i'): rPr.remove(e)
        rPr.insert(0, new_el('w:i'))
        chg = new_el('w:rPrChange', **{'w:id': nid(), 'w:author': AUTHOR, 'w:date': STAMP})
        inner = new_el('w:rPr')
        if old is not None:
            for c in list(old):
                if c.tag != W + 'rPrChange': inner.append(c)
        chg.append(inner); rPr.append(chg)

# ---------- an edit queue, applied in descending offset order ----------
class Edits:
    def __init__(self, p): self.p, self.q = p, []
    def delete(self, s, e):        self.q.append((s, 'del', (s, e)))
    def insert(self, at, txt, it=False): self.q.append((at, 'ins', (at, txt, it)))
    def hl(self, s, e, c="yellow"): self.q.append((s, 'hl', (s, e, c)))
    def ital(self, s, e):          self.q.append((s, 'it', (s, e)))
    def apply(self):
        # descending by offset; within same offset, formatting before text edits
        order = {'hl': 0, 'it': 1, 'del': 2, 'ins': 3}
        for _, kind, args in sorted(self.q, key=lambda x: (-x[0], order[x[1]])):
            if kind == 'del': track_delete(self.p, *args)
            elif kind == 'ins': track_insert(self.p, *args)
            elif kind == 'hl': highlight(self.p, *args)
            elif kind == 'it': track_italicize(self.p, *args)

# ---------- build sample ----------
doc = Document()
for pieces in [
    [("Adams, John, 15, 10, 16; and the Continental Congress, 22; see also Continental Congress", None)],
    [("Baltimore, Maryland,  30-20, 45n12, 82 and n, 128", None)],
    [("Churchill, Winston, 100-9, 429nn, 304nn1, 7", None)],
]:
    p = doc.add_paragraph()
    for text, it in pieces:
        r = p.add_run(text)
        if it is not None: r.italic = it
doc.save("sample2.docx")

# ---------- apply real index fixes ----------
doc = Document("sample2.docx"); ps = doc.paragraphs

t, _ = text_and_fmt(ps[0]._p)
e = Edits(ps[0]._p)
i = t.index("15, 10, 16"); e.hl(i, i + 10, "yellow")          # out of numeric order
j = t.index("see also");   e.ital(j, j + 8)                    # italicise see also
e.apply()

t, _ = text_and_fmt(ps[1]._p)
e = Edits(ps[1]._p)
k = t.index("30-20")
e.delete(k + 2, k + 3); e.insert(k + 2, "–")              # hyphen -> en dash
m = re.search(r"  ", t); e.delete(m.start(), m.start() + 1)     # double space
n = t.index("45n12"); e.ital(n + 2, n + 3)                      # italicise the n
e.hl(k, k + 5, "yellow")                                        # range 30-20 reversed
e.apply()

t, _ = text_and_fmt(ps[2]._p)
e = Edits(ps[2]._p)
q = t.index("100-9"); e.hl(q, q + 5, "cyan")                    # elided range
z = t.index("429nn"); e.ital(z + 3, z + 5)                      # italicise nn
e.apply()

doc.save("checked2.docx")
print("saved checked2.docx")
