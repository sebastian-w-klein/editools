"""Running the rules over a document and marking it up."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

from . import docxio, parser, rules
from .docxio import EditQueue, _Ids
from .parser import Entry


@dataclass
class Result:
    path: Path
    entries: int = 0
    findings: list[tuple[int, rules.Finding]] = field(default_factory=list)
    #: entry term by paragraph position, so a finding can name where it is
    terms: dict[int, str] = field(default_factory=dict)

    def counts(self) -> dict[str, int]:
        tally: dict[str, int] = {}
        for _, finding in self.findings:
            tally[finding.rule] = tally.get(finding.rule, 0) + 1
        return tally

    @property
    def fixes(self) -> int:
        return sum(1 for _, f in self.findings if f.action)

    @property
    def flags(self) -> int:
        return sum(1 for _, f in self.findings if not f.action)

    def described(self) -> list[tuple[str, rules.Finding]]:
        """Findings paired with the entry each one sits in, in document order."""
        return [(self.terms.get(position, ""), finding)
                for position, finding in sorted(self.findings,
                                                key=lambda x: x[0])]


def parse_document(document) -> list[tuple[int, Entry, object]]:
    """Parse every non-empty paragraph into an entry."""
    parsed = []
    for position, paragraph in enumerate(document.paragraphs):
        text, italics = docxio.read(paragraph._p)
        if not text.strip():
            continue
        parsed.append((position, parser.parse(text, italics), paragraph))
    return parsed


def _without_overlaps(findings: list[rules.Finding]) -> list[rules.Finding]:
    """Drop edits that overlap an edit already accepted for this paragraph.

    Two rules can land on the same characters — an elision fix and a dash fix
    inside one range, say. Applying both would corrupt the text, so the first
    wins and the second is dropped. Flags never overlap-conflict because they
    only add a highlight.
    """
    kept: list[rules.Finding] = []
    taken: list[tuple[int, int]] = []
    for finding in sorted(findings, key=lambda f: (f.action is None, f.start)):
        if finding.action:
            if any(finding.start < stop and start < finding.stop
                   for start, stop in taken):
                continue
            taken.append((finding.start, finding.stop))
        kept.append(finding)
    return kept


def _queue(queue: EditQueue, finding: rules.Finding) -> None:
    """Turn one finding into tracked edits."""
    if finding.action == "delete":
        queue.delete(finding.start, finding.stop)
    elif finding.action == "replace":
        queue.delete(finding.start, finding.stop)
        queue.insert(finding.start, finding.replacement, italic=finding.italic)
    elif finding.action == "italicise":
        queue.italicise(finding.start, finding.stop)
    else:
        queue.highlight(finding.start, finding.stop, finding.colour,
                        note=finding.message)


def run(path: str | Path, out_path: str | Path | None = None,
        mark: bool = True, last_page: int | None = None) -> Result:
    """Check an index and, unless told otherwise, write a marked-up copy."""
    path = Path(path)
    document = Document(str(path))
    parsed = parse_document(document)
    result = Result(path=path,
                    entries=sum(1 for _, e, _ in parsed if not e.is_note))
    result.terms = {position: entry.term for position, entry, _ in parsed}

    for position, entry, _ in parsed:
        if entry.is_note:
            continue
        found = []
        for rule in rules.PER_ENTRY:
            found.extend(rule(entry))
        found.extend(rules.check_page_too_high(entry, last_page))
        for finding in _without_overlaps(found):
            result.findings.append((position, finding))

    for position, finding in rules.check_entry_order(
            [(p, e) for p, e, _ in parsed]):
        result.findings.append((position, finding))

    if mark:
        ids = _Ids()
        by_position: dict[int, list[rules.Finding]] = {}
        for position, finding in result.findings:
            by_position.setdefault(position, []).append(finding)
        for position, findings in by_position.items():
            queue = EditQueue(document, document.paragraphs[position]._p, ids)
            for finding in findings:
                _queue(queue, finding)
            queue.apply()
        out_path = Path(out_path or path.with_name(path.stem + "_checked.docx"))
        document.save(str(out_path))
        result.path = out_path

    return result
