"""End to end: an index with known faults comes back marked up."""

import zipfile

from docx import Document

from editools.index import audit
from editools.index.docxio import W


def test_a_clean_index_is_left_alone(make_docx):
    path = make_docx(
        "Adams, John, 15, 22, 30",
        "Baltimore, 45, 47–52",
        "Chicago, 60; architecture, 61; parks, 64",
    )
    result = audit.run(path, mark=False)
    assert result.entries == 3
    assert result.findings == []


def test_the_four_faults_are_all_found(make_docx):
    path = make_docx(
        "Adams, John, 15, 10, 16",                     # page order
        "Baltimore, 30–20",                            # reversed range
        "Chicago, 60; parks, 61; architecture, 64",    # subentry order
        "Aachen, 70",                                  # entry order
    )
    result = audit.run(path, mark=False)
    assert set(result.counts()) == {
        "page-order", "range-order", "subentry-order", "entry-order"}


def test_findings_name_the_entry_they_are_in(make_docx):
    """The label keeps the whole name, though it sorts on 'Adams' alone."""
    path = make_docx("Adams, John, 15, 10, 16")
    result = audit.run(path, mark=False)
    assert result.described()[0][0] == "Adams, John"


def test_the_marked_up_file_carries_highlights_and_comments(make_docx, tmp_path):
    path = make_docx("Adams, John, 15, 10, 16")
    out = tmp_path / "checked.docx"
    audit.run(path, out_path=out)

    document = Document(str(out))
    paragraph = document.paragraphs[0]._p
    assert paragraph.findall(".//" + W + "highlight")
    # tracked, so rejecting the changes takes the highlight away again
    assert paragraph.findall(".//" + W + "rPrChange")
    with zipfile.ZipFile(str(out)) as archive:
        assert "word/comments.xml" in archive.namelist()
        assert "ascending order" in archive.read("word/comments.xml").decode()


def test_the_illustration_note_is_not_counted_as_an_entry(make_docx):
    path = make_docx(
        "Page numbers in italics refer to illustrations",
        "Adams, John, 15",
    )
    result = audit.run(path, mark=False)
    assert result.entries == 1
    assert result.findings == []


def test_fixes_are_applied_as_tracked_edits(make_docx):
    """A messy entry comes back corrected, and rejecting restores it."""
    from index.test_docxio import resolve

    messy = 'Adams, John,  15-22, 82n, 308–310; See also "Smith"'
    path = make_docx(messy)
    out = path.with_name("checked.docx")
    audit.run(path, out_path=out)

    paragraph = Document(str(out)).paragraphs[0]._p
    assert resolve(paragraph, "reject") == messy
    assert resolve(paragraph, "accept") == (
        'Adams, John, 15–22, 82n, 308–10; see also “Smith”')


def test_overlapping_fixes_do_not_corrupt_the_text(make_docx):
    """An elision fix and a dash fix can land on the same range."""
    from index.test_docxio import resolve

    messy = "Adams, 308-310, 100-9"
    path = make_docx(messy)
    out = path.with_name("checked.docx")
    audit.run(path, out_path=out)

    paragraph = Document(str(out)).paragraphs[0]._p
    assert resolve(paragraph, "reject") == messy
    accepted = resolve(paragraph, "accept")
    assert "–" in accepted and "-" not in accepted


def test_a_page_past_the_end_of_the_book_is_flagged(make_docx):
    path = make_docx("Adams, John, 15, 402")
    assert audit.run(path, mark=False, last_page=390).counts() == {
        "page-too-high": 1}
    assert audit.run(path, mark=False).counts() == {}
