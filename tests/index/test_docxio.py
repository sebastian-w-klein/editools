"""The tracked-changes layer.

The contract these tests hold to: accepting every change gives the corrected
text, and rejecting every change gives back exactly what came in.
"""

import zipfile

from docx import Document
from lxml import etree

from editools.index import docxio
from editools.index.docxio import EditQueue, W, _Ids

NS = {"w": W[1:-1]}


def resolve(paragraph, mode):
    """The paragraph's text as Word would show it after accept-all/reject-all."""
    out = []
    for run in paragraph.iter(W + "r"):
        ancestors = {a.tag for a in run.iterancestors()}
        deleted, inserted = W + "del" in ancestors, W + "ins" in ancestors
        if mode == "accept" and deleted:
            continue
        if mode == "reject" and inserted:
            continue
        node = run.find(W + "delText") if deleted else run.find(W + "t")
        if node is not None and node.text:
            out.append(node.text)
    return "".join(out)


def edit(path, fn):
    """Apply ``fn(queue, text)`` to the first paragraph and reload the file."""
    document = Document(str(path))
    paragraph = document.paragraphs[0]
    text, _ = docxio.read(paragraph._p)
    queue = EditQueue(document, paragraph._p, _Ids())
    fn(queue, text)
    queue.apply()
    document.save(str(path))
    return Document(str(path)).paragraphs[0]._p, text


def test_reading_text_and_italics(make_docx):
    path = make_docx([("McDuffie, 15, 82", None), ("n", True), (", 128", None)])
    text, italics = docxio.read(Document(str(path)).paragraphs[0]._p)
    assert text == "McDuffie, 15, 82n, 128"
    assert [t for t, i in zip(text, italics) if i] == ["n"]


def test_explicit_not_italic_is_not_italic(make_docx):
    """Word writes <w:i w:val="0"/> for roman; presence alone is not italic."""
    path = make_docx([("roman", False), ("italic", True)])
    text, italics = docxio.read(Document(str(path)).paragraphs[0]._p)
    assert italics == [False] * 5 + [True] * 6


def test_replacing_a_hyphen_with_an_en_dash(make_docx):
    path = make_docx("Baltimore, 30-20, 45")
    def fn(queue, text):
        at = text.index("30-20") + 2
        queue.delete(at, at + 1)
        queue.insert(at, "–")
    paragraph, before = edit(path, fn)
    assert resolve(paragraph, "accept") == "Baltimore, 30–20, 45"
    assert resolve(paragraph, "reject") == before


def test_several_edits_do_not_disturb_each_other(make_docx):
    """Edits are queued and applied back to front, so earlier offsets hold."""
    path = make_docx("Baltimore,  30-20, 45n12, 82")
    def fn(queue, text):
        queue.delete(text.index("  "), text.index("  ") + 1)   # double space
        at = text.index("30-20") + 2
        queue.delete(at, at + 1)                               # hyphen
        queue.insert(at, "–")                                  # en dash
        n = text.index("45n12") + 2
        queue.italicise(n, n + 1)                              # the note "n"
    paragraph, before = edit(path, fn)
    assert resolve(paragraph, "accept") == "Baltimore, 30–20, 45n12, 82"
    assert resolve(paragraph, "reject") == before


def test_a_highlight_is_tracked_so_rejecting_removes_it(make_docx):
    path = make_docx("Adams, 15, 10, 16")
    paragraph, _ = edit(path, lambda q, t: q.highlight(7, 17, "yellow"))
    assert paragraph.findall(".//" + W + "highlight")
    assert paragraph.findall(".//" + W + "rPrChange")


def test_italicising_marks_only_the_target(make_docx):
    path = make_docx("art, 290; see also films")
    paragraph, _ = edit(path, lambda q, t: q.italicise(t.index("see also"),
                                                       t.index("see also") + 8))
    italic = [r.find(W + "t").text for r in paragraph.iter(W + "r")
              if docxio.toggle_on(r.find(W + "rPr"), "i")]
    assert italic == ["see also"]


def test_a_comment_carries_the_reason(make_docx):
    path = make_docx("Adams, 15, 10, 16")
    document = Document(str(path))
    queue = EditQueue(document, document.paragraphs[0]._p, _Ids())
    queue.highlight(11, 13, "yellow", note="10 comes after 15")
    queue.apply()
    document.save(str(path))
    with zipfile.ZipFile(str(path)) as archive:
        comments = archive.read("word/comments.xml").decode()
    assert "10 comes after 15" in comments
    assert docxio.AUTHOR in comments


def test_the_file_still_opens_and_the_xml_is_valid(make_docx):
    path = make_docx("Adams, 15, 10, 16", "Baltimore, 30-20")
    edit(path, lambda q, t: q.highlight(0, 5, "cyan", note="check this"))
    with zipfile.ZipFile(str(path)) as archive:
        etree.fromstring(archive.read("word/document.xml"))
    assert len(Document(str(path)).paragraphs) == 2
